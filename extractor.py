import os
import io
import base64
import logging
from pathlib import Path
from dataclasses import dataclass
from typing import Optional

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# Data class for extracted content
# ─────────────────────────────────────────────

@dataclass
class ExtractedContent:
    file_name: str
    file_type: str          # "pdf", "image", "text", "docx", "spreadsheet", "unknown"
    text: str               # extracted text
    page_count: int = 1
    used_ocr: bool = False
    used_vision: bool = False
    error: Optional[str] = None


# ─────────────────────────────────────────────
# MIME type → category mapping
# ─────────────────────────────────────────────

MIME_MAP = {
    # PDFs
    "application/pdf": "pdf",

    # Images
    "image/jpeg": "image",
    "image/png":  "image",
    "image/gif":  "image",
    "image/webp": "image",
    "image/bmp":  "image",
    "image/tiff": "image",

    # Plain text
    "text/plain": "text",

    # Word / Google Docs
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
    "application/vnd.google-apps.document": "gdoc",

    # Spreadsheets
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "spreadsheet",
    "application/vnd.ms-excel": "spreadsheet",
    "text/csv": "spreadsheet",
    "application/vnd.google-apps.spreadsheet": "gsheet",
}


# ─────────────────────────────────────────────
# PDF Extractor
# ─────────────────────────────────────────────

def extract_pdf(file_bytes: bytes, file_name: str) -> ExtractedContent:
    """Extract text from PDF using PyMuPDF. Falls back to OCR if text is sparse."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text())

        full_text = "\n\n".join(pages_text).strip()

        # If very little text extracted, the PDF is likely scanned → OCR
        if len(full_text) < 100:
            logger.info(f"[PDF] Sparse text in '{file_name}', switching to OCR...")
            return extract_pdf_with_ocr(doc, file_name)

        return ExtractedContent(
            file_name=file_name,
            file_type="pdf",
            text=full_text,
            page_count=len(doc),
        )

    except ImportError:
        return ExtractedContent(
            file_name=file_name, file_type="pdf", text="",
            error="PyMuPDF not installed. Run: pip install pymupdf"
        )
    except Exception as e:
        return ExtractedContent(file_name=file_name, file_type="pdf", text="", error=str(e))


def extract_pdf_with_ocr(doc, file_name: str) -> ExtractedContent:
    """Render each PDF page as image → OCR with pytesseract."""
    try:
        import pytesseract
        from PIL import Image
        import fitz

        pages_text = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img)
            pages_text.append(text)

        return ExtractedContent(
            file_name=file_name,
            file_type="pdf",
            text="\n\n".join(pages_text).strip(),
            page_count=len(doc),
            used_ocr=True,
        )

    except ImportError:
        return ExtractedContent(
            file_name=file_name, file_type="pdf", text="",
            error="pytesseract or Pillow not installed for OCR fallback."
        )


# ─────────────────────────────────────────────
# Image Extractor  (OCR  +  Groq Vision)
# ─────────────────────────────────────────────

def extract_image_ocr(file_bytes: bytes, file_name: str) -> ExtractedContent:
    """Extract text from image using pytesseract OCR."""
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(img).strip()

        return ExtractedContent(
            file_name=file_name, file_type="image",
            text=text or "[No text found by OCR]",
            used_ocr=True,
        )
    except ImportError:
        return ExtractedContent(
            file_name=file_name, file_type="image", text="",
            error="pytesseract / Pillow not installed. Run: pip install pytesseract pillow"
        )
    except Exception as e:
        return ExtractedContent(file_name=file_name, file_type="image", text="", error=str(e))


def extract_image_vision(file_bytes: bytes, file_name: str, mime_type: str) -> ExtractedContent:
    """
    Describe / extract information from an image using Groq's vision model.
    Uses llama-4-scout-17b-16e-instruct which supports vision.
    """
    try:
        from groq import Groq

        groq_api_key = os.getenv("GROQ_API_KEY")
        if not groq_api_key:
            raise ValueError("GROQ_API_KEY not set in environment.")

        client = Groq(api_key=groq_api_key)

        # Encode image to base64
        b64_image = base64.b64encode(file_bytes).decode("utf-8")
        data_url = f"data:{mime_type};base64,{b64_image}"

        response = client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": data_url},
                        },
                        {
                            "type": "text",
                            "text": (
                                "Please extract and describe ALL text, data, tables, "
                                "charts, or meaningful content visible in this image. "
                                "Be thorough and structured."
                            ),
                        },
                    ],
                }
            ],
            max_tokens=1024,
        )

        description = response.choices[0].message.content.strip()

        return ExtractedContent(
            file_name=file_name, file_type="image",
            text=description,
            used_vision=True,
        )

    except ImportError:
        return ExtractedContent(
            file_name=file_name, file_type="image", text="",
            error="groq package not installed. Run: pip install groq"
        )
    except Exception as e:
        return ExtractedContent(file_name=file_name, file_type="image", text="", error=str(e))


# ─────────────────────────────────────────────
# Plain Text Extractor
# ─────────────────────────────────────────────

def extract_text(file_bytes: bytes, file_name: str) -> ExtractedContent:
    """Decode plain text file."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            text = file_bytes.decode(encoding).strip()
            return ExtractedContent(file_name=file_name, file_type="text", text=text)
        except UnicodeDecodeError:
            continue
    return ExtractedContent(
        file_name=file_name, file_type="text", text="",
        error="Could not decode text file with utf-8 / latin-1 / cp1252."
    )


# ─────────────────────────────────────────────
# DOCX Extractor
# ─────────────────────────────────────────────

def extract_docx(file_bytes: bytes, file_name: str) -> ExtractedContent:
    """Extract text from .docx using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        full_text = "\n".join(paragraphs)
        if table_texts:
            full_text += "\n\n[Tables]\n" + "\n".join(table_texts)

        return ExtractedContent(file_name=file_name, file_type="docx", text=full_text.strip())

    except ImportError:
        return ExtractedContent(
            file_name=file_name, file_type="docx", text="",
            error="python-docx not installed. Run: pip install python-docx"
        )
    except Exception as e:
        return ExtractedContent(file_name=file_name, file_type="docx", text="", error=str(e))


# ─────────────────────────────────────────────
# Spreadsheet Extractor (XLSX / CSV)
# ─────────────────────────────────────────────

def extract_spreadsheet(file_bytes: bytes, file_name: str, mime_type: str) -> ExtractedContent:
    """Convert spreadsheet to readable text using pandas."""
    try:
        import pandas as pd

        if mime_type == "text/csv" or file_name.endswith(".csv"):
            df_dict = {"Sheet1": pd.read_csv(io.BytesIO(file_bytes))}
        else:
            xl = pd.ExcelFile(io.BytesIO(file_bytes))
            df_dict = {sheet: xl.parse(sheet) for sheet in xl.sheet_names}

        text_parts = []
        for sheet_name, df in df_dict.items():
            text_parts.append(f"[Sheet: {sheet_name}]")
            text_parts.append(df.to_string(index=False))

        return ExtractedContent(
            file_name=file_name, file_type="spreadsheet",
            text="\n\n".join(text_parts).strip()
        )

    except ImportError:
        return ExtractedContent(
            file_name=file_name, file_type="spreadsheet", text="",
            error="pandas / openpyxl not installed. Run: pip install pandas openpyxl"
        )
    except Exception as e:
        return ExtractedContent(file_name=file_name, file_type="spreadsheet", text="", error=str(e))


# ─────────────────────────────────────────────
# Master Dispatcher
# ─────────────────────────────────────────────

def extract_content(
    file_bytes: bytes,
    file_name: str,
    mime_type: str,
    use_vision_for_images: bool = True,   # True → Groq Vision; False → pytesseract OCR
) -> ExtractedContent:
    """
    Master extractor. Routes to the correct extractor based on MIME type.
    """
    category = MIME_MAP.get(mime_type, "unknown")

    # Fallback: detect from file extension if MIME is unknown
    if category == "unknown":
        ext = Path(file_name).suffix.lower()
        ext_map = {
            ".pdf": "pdf", ".txt": "text", ".md": "text",
            ".docx": "docx", ".doc": "docx",
            ".xlsx": "spreadsheet", ".xls": "spreadsheet", ".csv": "spreadsheet",
            ".jpg": "image", ".jpeg": "image", ".png": "image",
            ".gif": "image", ".bmp": "image", ".tiff": "image", ".webp": "image",
        }
        category = ext_map.get(ext, "unknown")

    logger.info(f"Extracting '{file_name}' | MIME: {mime_type} | Category: {category}")

    if category == "pdf":
        return extract_pdf(file_bytes, file_name)

    elif category == "image":
        if use_vision_for_images:
            result = extract_image_vision(file_bytes, file_name, mime_type)
            # If vision fails, fallback to OCR
            if result.error:
                logger.warning(f"Vision failed for '{file_name}', falling back to OCR. Error: {result.error}")
                return extract_image_ocr(file_bytes, file_name)
            return result
        else:
            return extract_image_ocr(file_bytes, file_name)

    elif category == "text":
        return extract_text(file_bytes, file_name)

    elif category in ("docx", "gdoc"):
        return extract_docx(file_bytes, file_name)

    elif category in ("spreadsheet", "gsheet"):
        return extract_spreadsheet(file_bytes, file_name, mime_type)

    else:
        return ExtractedContent(
            file_name=file_name, file_type="unknown", text="",
            error=f"Unsupported file type: {mime_type}"
        )
